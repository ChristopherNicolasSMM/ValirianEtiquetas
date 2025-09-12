import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
from copy import deepcopy
import sys
#from docx.oxml import OxmlElement

class WordEtiquetaHandler:
    def __init__(self, template_path: str):
        self.template_path = template_path
        self.output_dir = os.path.join(os.path.dirname(template_path), 'output')
        os.makedirs(self.output_dir, exist_ok=True)

    def _limpar_celula(self, cell):
        """Remove todo o conteúdo da célula (parágrafos e tabelas internas)."""
        for p in cell.paragraphs:
            if hasattr(p, "clear"):
                p.clear()
            else:
                p._element.clear_content()

        for t in list(cell.tables):
            t._element.getparent().remove(t._element)
            
        
    def _preencher_tabela(self, modelo_tabela, cell, dados):
        # Faz uma cópia profunda do XML da tabela modelo
        nova_tabela_element = deepcopy(modelo_tabela._element)

        # Anexa essa cópia ao XML da célula de destino
        cell._element.append(nova_tabela_element)

        # Recupera a tabela recém clonada (última da célula)
        nova_tabela = cell.tables[-1]
                
        # Substitui os placeholders pelos valores
        for i, row in enumerate(nova_tabela.rows):
            for j, tgt_cell in enumerate(row.cells):
                for p in tgt_cell.paragraphs:
                    for key, value in dados.items():
                        if f'{{{key}}}' in p.text:
                                                              
                            # Cria um novo run com o texto substituído
                            if key == "receita":                                
                                novo_texto =  p.text.replace(f'{{{key}}}', str(value))
                                p.text = ''
                                # Limpa os runs existentes
                                for r in p.runs:
                                    r.text = ""                                
                                run = p.add_run(novo_texto)
                                run.font.size = Pt(14)   # tamanho de fonte em pontos                            
                            else:
                                p.text = p.text.replace(f'{{{key}}}', str(value))
                                
                                
        #Remove visibilidade das bordas
        nova_tabela.style.style_id = 'None'
        nova_tabela.style.hidden   = False


    def _calcular_etiquetas_por_pagina(self) -> int:
        """Calcula quantas etiquetas cabem em uma página do template (colunas ímpares)."""
        doc = Document(self.template_path)
        tabela = doc.tables[0]
        etiquetas_por_pagina = 0

        for row in tabela.rows:
            for col_idx in range(len(row.cells)):
                if col_idx % 2 == 0:  # só colunas ímpares (0,2,4,...)
                    etiquetas_por_pagina += 1

        return etiquetas_por_pagina

    def criar_etiquetas(self, dados_lote: dict, quantidade_total: int, pagina: int = 1, extra_tags: dict | None = None):
        """Cria etiquetas em uma única página, respeitando o limite."""
        doc = Document(self.template_path)

        tabela_principal = doc.tables[0]
        modelo_tabela = tabela_principal.cell(0, 0).tables[0]

        dados = {
            'lote': dados_lote.get('batchNo', ''),
            'receita': dados_lote.get('name', ''),
            'abv': f"{dados_lote.get('measuredAbv', '')}%" if dados_lote.get('measuredAbv') else '',
            'ibu': str(dados_lote.get('estimatedIbu', '')),
            'estimatedColor': str(dados_lote.get('estimatedColor', '')),
            'data_brassagem': dados_lote.get('brewDate', ''),
            'data_engarrafamento': dados_lote.get('bottling_event', {}).get('time', '') if dados_lote.get('bottling_event') else '',
            'data_impressao': datetime.now().strftime('%d/%m/%Y %H:%M')
        }

        if extra_tags:
            for k, v in extra_tags.items():
                if k not in dados:
                    dados[k] = v

        etiquetas_preenchidas = 0

        for row in tabela_principal.rows:
            for col_idx, cell in enumerate(row.cells):
                if etiquetas_preenchidas >= quantidade_total:
                    break

                if col_idx % 2 == 1:  # pular colunas pares
                    continue

                self._limpar_celula(cell)
                self._preencher_tabela(modelo_tabela, cell, dados)
                etiquetas_preenchidas += 1

            if etiquetas_preenchidas >= quantidade_total:
                break

        nome_arquivo = f"etiqueta_{dados_lote['batchNo']}_p{pagina}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        caminho_saida = os.path.join(self.output_dir, nome_arquivo)
        doc.save(caminho_saida)
        return caminho_saida

    def criar_multiplas_paginas(self, dados_lote: dict, quantidade_total: int, extra_tags: dict | None = None):
        """Divide a geração de etiquetas em múltiplas páginas se necessário."""
        etiquetas_por_pagina = self._calcular_etiquetas_por_pagina()
        paginas_necessarias = (quantidade_total + etiquetas_por_pagina - 1) // etiquetas_por_pagina

        arquivos = []
        restantes = quantidade_total

        for pagina in range(1, paginas_necessarias + 1):
            qtd_nesta_pagina = min(restantes, etiquetas_por_pagina)
            arquivo = self.criar_etiquetas(dados_lote, qtd_nesta_pagina, pagina, extra_tags=extra_tags)
            arquivos.append(arquivo)
            restantes -= qtd_nesta_pagina


        return arquivos
